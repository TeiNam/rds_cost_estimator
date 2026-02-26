"""
문서 파일 텍스트 추출 및 Bedrock 파싱 모듈.

이 모듈은 PDF, DOCX, TXT 파일에서 텍스트를 추출하고
AWS Bedrock(Claude 모델)을 호출하여 인스턴스 사양 정보를 파싱하는
DocumentParser 클래스를 제공합니다.
"""

from __future__ import annotations

import logging
import os
from typing import TYPE_CHECKING, Literal

from rds_cost_estimator.exceptions import UnsupportedFileFormatError
from rds_cost_estimator.models import ParsedDocumentInfo

if TYPE_CHECKING:
    from rds_cost_estimator.bedrock_client import BedrockClient

# 모듈 레벨 로거 설정
logger = logging.getLogger(__name__)

# 지원하는 파일 형식 목록
SUPPORTED_FORMATS: list[str] = [".pdf", ".docx", ".txt", ".md", ".out"]


class DocumentParser:
    """문서 파일에서 텍스트를 추출하고 Bedrock으로 인스턴스 사양 정보를 파싱하는 클래스.

    단일 파일 또는 디렉토리를 입력받을 수 있습니다.
    디렉토리 입력 시 폴더 내 모든 지원 파일(PDF/DOCX/TXT/MD)의 텍스트를
    합쳐서 Bedrock에 전달합니다.

    Attributes:
        _bedrock_client: Bedrock API 호출을 담당하는 클라이언트 인스턴스
    """

    def __init__(self, bedrock_client: "BedrockClient") -> None:
        """DocumentParser 초기화.

        Args:
            bedrock_client: AWS Bedrock Runtime 클라이언트 인스턴스
        """
        # Bedrock 클라이언트 저장
        self._bedrock_client = bedrock_client

    def parse(self, file_path: str) -> ParsedDocumentInfo:
        """파일 또는 디렉토리를 입력받아 텍스트를 추출한 뒤 Bedrock으로 파싱.

        AWR .out 파일은 직접 파싱하여 Bedrock 컨텍스트에서 제외합니다.
        직접 파싱 결과를 먼저 적용하고, Bedrock 결과는 누락 필드만 보완합니다.

        Args:
            file_path: 파싱할 문서 파일 또는 디렉토리 경로

        Returns:
            문서에서 추출한 인스턴스 사양 정보

        Raises:
            UnsupportedFileFormatError: 지원하지 않는 파일 형식인 경우
            DocumentParseError: Bedrock API 호출 실패 또는 응답 파싱 실패 시
        """
        logger.info("문서 파싱 시작: %s", file_path)

        # 1단계: AWR .out 파일 직접 파싱 (Bedrock 호출 전)
        awr_parsed = self._parse_awr_out_full(file_path)

        # 2단계: .out 제외한 텍스트 추출 → Bedrock 호출
        if os.path.isdir(file_path):
            text = self._extract_text_from_directory(file_path)
        else:
            text = self._extract_text(file_path)

        logger.debug("텍스트 추출 완료 (.out 제외): %d자", len(text))

        # Bedrock 호출 (텍스트가 비어있으면 빈 ParsedDocumentInfo 반환)
        if text.strip():
            result = self._bedrock_client.invoke(text)
            logger.info("Bedrock 파싱 완료: %s", file_path)
        else:
            result = ParsedDocumentInfo()
            logger.info("Bedrock 컨텍스트 없음, 직접 파싱 결과만 사용")

        # 3단계: 직접 파싱 결과를 Bedrock 결과에 병합 (직접 파싱 우선)
        self._apply_awr_parsed(awr_parsed, result)

        # 4단계: migration_recommendation.md에서 타겟 엔진 직접 파싱
        self._supplement_target_engine(file_path, result)

        return result

    def _extract_text_from_directory(self, dir_path: str) -> str:
        """디렉토리 내 지원 파일에서 텍스트를 추출하여 합칩니다.

        .out 파일은 직접 파싱으로 대체하므로 Bedrock 컨텍스트에서 제외합니다.
        PDF/DOCX/TXT/MD 파일만 Bedrock에 전달합니다.

        Args:
            dir_path: 탐색할 디렉토리 경로

        Returns:
            모든 파일의 텍스트를 합친 문자열

        Raises:
            UnsupportedFileFormatError: 디렉토리에 지원 파일이 없는 경우
        """
        logger.info("디렉토리 내 파일 탐색 시작: %s", dir_path)

        # 지원 파일 확장자 필터링 (.out 제외 - 직접 파싱으로 대체)
        supported_files: list[str] = []
        has_any_supported = False
        for entry in sorted(os.listdir(dir_path)):
            entry_path = os.path.join(dir_path, entry)
            if not os.path.isfile(entry_path):
                continue
            dot_index = entry.rfind(".")
            if dot_index == -1:
                continue
            ext = entry[dot_index:].lower()
            if ext in (".pdf", ".docx", ".txt", ".md", ".out"):
                has_any_supported = True
                # .out 파일은 직접 파싱으로 대체 → Bedrock 컨텍스트에서 제외
                if ext != ".out":
                    supported_files.append(entry_path)

        if not has_any_supported:
            logger.warning("디렉토리에 지원 파일 없음: %s", dir_path)
            raise UnsupportedFileFormatError(dir_path, SUPPORTED_FORMATS)

        if not supported_files:
            # .out 파일만 있는 경우 빈 텍스트 반환 (직접 파싱으로 처리)
            logger.info("디렉토리에 .out 파일만 존재, Bedrock 컨텍스트 비어있음")
            return ""

        logger.info("디렉토리에서 %d개 파일 발견 (.out 제외): %s",
                     len(supported_files),
                     [os.path.basename(f) for f in supported_files])

        # 각 파일의 텍스트를 구분자와 함께 결합
        text_parts: list[str] = []
        for fpath in supported_files:
            fname = os.path.basename(fpath)
            logger.debug("파일 텍스트 추출: %s", fname)
            file_text = self._extract_text(fpath)
            text_parts.append(
                f"=== 파일: {fname} ===\n{file_text}"
            )

        combined = "\n\n".join(text_parts)
        logger.info("디렉토리 텍스트 합치기 완료: %d개 파일, 총 %d자",
                     len(supported_files), len(combined))
        return combined

    def _extract_text(self, file_path: str) -> str:
        """파일 형식별 텍스트 추출.

        파일 형식을 감지하여 적절한 라이브러리로 텍스트를 추출합니다:
        - PDF: pypdf.PdfReader로 페이지별 텍스트 추출 후 결합
        - DOCX: docx.Document로 단락(paragraph) 텍스트 추출 후 결합
        - TXT: 내장 open()으로 UTF-8 직접 읽기

        Args:
            file_path: 텍스트를 추출할 파일 경로

        Returns:
            추출된 텍스트 문자열

        Raises:
            UnsupportedFileFormatError: 지원하지 않는 파일 형식인 경우
        """
        # 파일 형식 감지
        fmt = self._detect_format(file_path)

        if fmt == "pdf":
            # PDF: pypdf 라이브러리로 페이지별 텍스트 추출
            return self._extract_text_from_pdf(file_path)
        elif fmt == "docx":
            # DOCX: python-docx 라이브러리로 단락 텍스트 추출
            return self._extract_text_from_docx(file_path)
        elif fmt == "md":
            # MD: 내장 open()으로 UTF-8 직접 읽기 (TXT와 동일)
            return self._extract_text_from_txt(file_path)
        else:
            # TXT: 내장 open()으로 UTF-8 직접 읽기
            return self._extract_text_from_txt(file_path)

    def _detect_format(self, file_path: str) -> Literal["pdf", "docx", "txt", "md"]:
        """파일 확장자로 형식 감지.

        파일 경로에서 확장자를 추출하여 지원하는 형식인지 확인합니다.
        확장자는 소문자로 변환하여 비교합니다.

        Args:
            file_path: 형식을 감지할 파일 경로

        Returns:
            감지된 파일 형식 ("pdf", "docx", "txt" 중 하나)

        Raises:
            UnsupportedFileFormatError: 지원하지 않는 파일 형식인 경우
        """
        # 파일 경로에서 확장자 추출 (소문자 변환)
        # 예: "/path/to/file.PDF" → ".pdf"
        dot_index = file_path.rfind(".")
        if dot_index == -1:
            # 확장자가 없는 경우
            ext = ""
        else:
            ext = file_path[dot_index:].lower()

        # 지원하는 형식 매핑
        if ext == ".pdf":
            return "pdf"
        elif ext == ".docx":
            return "docx"
        elif ext == ".txt":
            return "txt"
        elif ext == ".md":
            return "md"
        elif ext == ".out":
            # AWR .out 파일은 텍스트 형식으로 처리
            return "txt"
        else:
            # 지원하지 않는 형식이면 예외 발생
            logger.warning("지원하지 않는 파일 형식: %s (확장자: %s)", file_path, ext)
            raise UnsupportedFileFormatError(file_path, SUPPORTED_FORMATS)

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """PDF 파일에서 텍스트 추출.

        pypdf.PdfReader를 사용하여 각 페이지의 텍스트를 추출하고
        줄바꿈으로 결합합니다.

        Args:
            file_path: PDF 파일 경로

        Returns:
            추출된 텍스트 (페이지별 텍스트를 "\n"으로 결합)
        """
        import pypdf  # PDF 텍스트 추출 라이브러리

        logger.debug("PDF 텍스트 추출 시작: %s", file_path)

        # PdfReader로 PDF 파일 열기
        reader = pypdf.PdfReader(file_path)

        # 각 페이지에서 텍스트 추출
        page_texts: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                page_texts.append(page_text)

        # 페이지별 텍스트를 줄바꿈으로 결합
        result = "\n".join(page_texts)
        logger.debug("PDF 텍스트 추출 완료: %d페이지, %d자", len(reader.pages), len(result))

        return result

    def _extract_text_from_docx(self, file_path: str) -> str:
        """DOCX 파일에서 텍스트 추출.

        python-docx 라이브러리를 사용하여 각 단락(paragraph)의 텍스트를 추출하고
        줄바꿈으로 결합합니다.

        Args:
            file_path: DOCX 파일 경로

        Returns:
            추출된 텍스트 (단락별 텍스트를 "\n"으로 결합)
        """
        import docx  # python-docx 라이브러리 (Word 문서 처리)

        logger.debug("DOCX 텍스트 추출 시작: %s", file_path)

        # Document 객체로 DOCX 파일 열기
        document = docx.Document(file_path)

        # 각 단락에서 텍스트 추출
        paragraph_texts: list[str] = [
            paragraph.text for paragraph in document.paragraphs
        ]

        # 단락별 텍스트를 줄바꿈으로 결합
        result = "\n".join(paragraph_texts)
        logger.debug(
            "DOCX 텍스트 추출 완료: %d단락, %d자",
            len(document.paragraphs),
            len(result),
        )

        return result

    def _extract_text_from_txt(self, file_path: str) -> str:
        """TXT 파일에서 텍스트 추출.

        내장 open() 함수를 사용하여 UTF-8 인코딩으로 파일을 직접 읽습니다.

        Args:
            file_path: TXT 파일 경로

        Returns:
            파일 전체 내용 문자열
        """
        logger.debug("TXT 텍스트 추출 시작: %s", file_path)

        # UTF-8 인코딩으로 텍스트 파일 직접 읽기
        with open(file_path, "r", encoding="utf-8") as f:
            result = f.read()

        logger.debug("TXT 텍스트 추출 완료: %d자", len(result))

        return result

    def _parse_awr_out_full(self, file_path: str) -> dict:
        """AWR .out 파일의 모든 섹션을 직접 파싱합니다.

        OS-INFORMATION, MAIN-METRICS, MEMORY, SGA-ADVICE, PERCENT-CPU,
        SYSSTAT 섹션에서 ParsedDocumentInfo에 필요한 필드를 추출합니다.

        Returns:
            직접 파싱된 필드 딕셔너리 (키: ParsedDocumentInfo 필드명)
        """
        out_files = self._find_awr_out_files(file_path)
        if not out_files:
            return {}

        merged: dict = {}
        for out_file in out_files:
            logger.info("AWR .out 파일 전체 직접 파싱: %s", out_file)
            try:
                with open(out_file, "r", encoding="utf-8") as f:
                    content = f.read()
            except Exception as e:
                logger.warning("AWR .out 파일 읽기 실패: %s - %s", out_file, e)
                continue

            # 각 섹션 파싱
            os_info = self._parse_os_information(content)
            main_metrics = self._parse_main_metrics_full(content)
            memory = self._parse_memory_section(content)
            sga_advice = self._parse_sga_advice(content)
            network = self._parse_awr_out_network_from_content(content)

            # 결과 병합 (첫 번째 파일 우선)
            for d in [os_info, main_metrics, memory, sga_advice, network]:
                for k, v in d.items():
                    if v is not None and k not in merged:
                        merged[k] = v

        return merged

    def _apply_awr_parsed(self, awr: dict, parsed: ParsedDocumentInfo) -> None:
        """직접 파싱 결과를 ParsedDocumentInfo에 적용합니다.

        직접 파싱 결과가 있으면 Bedrock 결과를 덮어씁니다 (직접 파싱 우선).
        Bedrock에서만 추출 가능한 필드(recommended_instance 등)는 유지합니다.
        """
        if not awr:
            return

        # 기본 정보 (직접 파싱 우선)
        if awr.get("db_name"):
            parsed.db_name = awr["db_name"]
        if awr.get("oracle_version"):
            parsed.oracle_version = awr["oracle_version"]
        if awr.get("cpu_cores") is not None:
            parsed.cpu_cores = awr["cpu_cores"]
        if awr.get("num_cpus") is not None:
            parsed.num_cpus = awr["num_cpus"]
        if awr.get("physical_memory_gb") is not None:
            parsed.physical_memory_gb = awr["physical_memory_gb"]
        if awr.get("db_size_gb") is not None:
            parsed.db_size_gb = awr["db_size_gb"]
        if awr.get("instance_config"):
            parsed.instance_config = awr["instance_config"]

        # AWR 메트릭 (직접 파싱 우선 - 덮어쓰기)
        m = parsed.awr_metrics
        if awr.get("avg_cpu_percent") is not None:
            m.avg_cpu_percent = awr["avg_cpu_percent"]
        if awr.get("peak_cpu_percent") is not None:
            m.peak_cpu_percent = awr["peak_cpu_percent"]
        if awr.get("avg_cpu_per_s") is not None:
            m.avg_cpu_per_s = awr["avg_cpu_per_s"]
        if awr.get("peak_cpu_per_s") is not None:
            m.peak_cpu_per_s = awr["peak_cpu_per_s"]
        if awr.get("avg_iops") is not None:
            m.avg_iops = awr["avg_iops"]
        if awr.get("peak_iops") is not None:
            m.peak_iops = awr["peak_iops"]
        if awr.get("avg_memory_gb") is not None:
            m.avg_memory_gb = awr["avg_memory_gb"]
        if awr.get("peak_memory_gb") is not None:
            m.peak_memory_gb = awr["peak_memory_gb"]

        # 네트워크/Redo (직접 파싱 우선 - 덮어쓰기)
        if awr.get("sent_bytes_per_day") is not None:
            m.sqlnet_bytes_sent_per_day = awr["sent_bytes_per_day"]
        if awr.get("recv_bytes_per_day") is not None:
            m.sqlnet_bytes_received_per_day = awr["recv_bytes_per_day"]
        if awr.get("redo_bytes_per_day") is not None:
            m.redo_bytes_per_day = awr["redo_bytes_per_day"]

        # SGA 분석 (직접 파싱 우선)
        sga = parsed.sga_analysis
        if awr.get("current_sga_gb") is not None:
            sga.current_sga_gb = awr["current_sga_gb"]
        if awr.get("recommended_sga_gb") is not None:
            sga.recommended_sga_gb = awr["recommended_sga_gb"]

        logger.info("AWR .out 직접 파싱 결과 적용 완료: %d개 필드", len(awr))

    def _parse_os_information(self, content: str) -> dict:
        """AWR .out의 OS-INFORMATION 섹션에서 서버 기본 정보를 파싱합니다.

        추출 필드: db_name, oracle_version, cpu_cores, num_cpus,
                   physical_memory_gb, db_size_gb, instance_config, sga_target
        """
        import re

        match = re.search(
            r"~~BEGIN-OS-INFORMATION~~\s*\n(.*?)~~END-OS-INFORMATION~~",
            content, re.DOTALL,
        )
        if not match:
            return {}

        section = match.group(1)
        result: dict = {}

        # STAT_NAME → STAT_VALUE 매핑 파싱
        kv: dict[str, str] = {}
        for line in section.split("\n"):
            line = line.strip()
            if not line or line.startswith("---") or line.startswith("STAT_NAME"):
                continue
            # 고정 폭 형식: 이름(60자) + 값
            parts = line.split()
            if len(parts) >= 2:
                key = parts[0]
                value = " ".join(parts[1:])
                kv[key] = value

        # 필드 추출
        if "DB_NAME" in kv:
            result["db_name"] = kv["DB_NAME"]

        if "VERSION" in kv:
            result["oracle_version"] = kv["VERSION"]
        elif "BANNER" in kv:
            # BANNER에서 버전 추출: "Oracle Database 19c ... Release 19.0.0.0.0"
            banner = kv["BANNER"]
            ver_match = re.search(r"Release\s+([\d.]+)", banner)
            if ver_match:
                result["oracle_version"] = ver_match.group(1)

        if "NUM_CPU_CORES" in kv:
            try:
                result["cpu_cores"] = int(kv["NUM_CPU_CORES"])
            except ValueError:
                pass

        if "NUM_CPUS" in kv:
            try:
                result["num_cpus"] = int(kv["NUM_CPUS"])
            except ValueError:
                pass

        if "PHYSICAL_MEMORY_GB" in kv:
            try:
                result["physical_memory_gb"] = float(kv["PHYSICAL_MEMORY_GB"])
            except ValueError:
                pass

        if "TOTAL_DB_SIZE_GB" in kv:
            try:
                result["db_size_gb"] = float(kv["TOTAL_DB_SIZE_GB"])
            except ValueError:
                pass

        # 인스턴스 구성: "2 (RAC)" 형태
        if "INSTANCES" in kv:
            instances = kv["INSTANCES"]
            try:
                inst_count = int(instances)
                if inst_count > 1:
                    result["instance_config"] = f"{inst_count} (RAC)"
                else:
                    result["instance_config"] = "1 (Single)"
            except ValueError:
                result["instance_config"] = instances

        # SGA_TARGET (바이트 → GB)
        if "SGA_TARGET" in kv:
            try:
                sga_bytes = int(kv["SGA_TARGET"])
                if sga_bytes > 0:
                    result["current_sga_gb"] = round(sga_bytes / (1024 ** 3), 1)
            except ValueError:
                pass

        return result

    def _parse_main_metrics_full(self, content: str) -> dict:
        """AWR .out의 MAIN-METRICS 섹션에서 성능 메트릭을 파싱합니다.

        추출 필드: avg/peak CPU%, avg/peak CPU/s, avg/peak IOPS, redo_mb_s
        RAC 환경에서는 인스턴스별 합산(IOPS) 또는 평균(CPU%)을 계산합니다.

        주의: 'end' 컬럼은 날짜+시간(예: "26/01/15 09:00")으로 공백을 포함하여
        split() 시 2개 토큰이 됩니다. end 이후 컬럼만 +1 오프셋을 적용합니다.
        """
        import re

        match = re.search(
            r"~~BEGIN-MAIN-METRICS~~\s*\n(.*?)~~END-MAIN-METRICS~~",
            content, re.DOTALL,
        )
        if not match:
            return {}

        section = match.group(1)
        lines = section.strip().split("\n")

        # 헤더 행 찾기
        header_line = None
        data_start = 0
        for i, line in enumerate(lines):
            stripped = line.strip()
            if stripped.startswith("snap") or "os_cpu" in stripped:
                header_line = stripped
                data_start = i + 1
                break

        if not header_line:
            return {}

        headers = header_line.split()
        col_map = {h: idx for idx, h in enumerate(headers)}

        # 'end' 컬럼 위치 파악 (날짜+시간이 2개 토큰으로 분리됨)
        end_header_idx = col_map.get("end")

        # 구분선 건너뛰기
        for i in range(data_start, len(lines)):
            if lines[i].strip().startswith("---"):
                data_start = i + 1
                break

        def _col(cols: list[str], header_idx: int | None) -> str | None:
            """헤더 인덱스에서 end 컬럼 오프셋을 보정하여 데이터 값을 반환합니다.

            end 컬럼 이전: 오프셋 0 (snap, dur_m)
            end 컬럼 이후: 오프셋 +1 (inst, os_cpu, cpu_per_s, ...)
            """
            if header_idx is None:
                return None
            actual = header_idx
            if end_header_idx is not None and header_idx > end_header_idx:
                actual += 1
            if actual < len(cols):
                return cols[actual]
            return None

        # 필요한 컬럼 인덱스 (헤더 기준)
        os_cpu_idx = col_map.get("os_cpu")
        os_cpu_max_idx = col_map.get("os_cpu_max")
        cpu_per_s_idx = col_map.get("cpu_per_s")
        cpu_per_s_max_idx = col_map.get("cpu_per_s_max")
        read_iops_idx = col_map.get("read_iops")
        write_iops_idx = col_map.get("write_iops")
        read_iops_max_idx = col_map.get("read_iops_max")
        write_iops_max_idx = col_map.get("write_iops_max")
        redo_idx = col_map.get("redo_mb_s")
        dur_idx = col_map.get("dur_m")
        snap_idx = col_map.get("snap")
        inst_idx = col_map.get("inst")

        # 데이터 행 파싱 - 스냅샷×인스턴스별 수집
        rows: list[dict] = []
        for i in range(data_start, len(lines)):
            line = lines[i].strip()
            if not line or line.startswith("~~"):
                break
            cols = line.split()
            row: dict = {}
            try:
                v = _col(cols, snap_idx)
                if v is not None:
                    row["snap"] = v
                v = _col(cols, inst_idx)
                if v is not None:
                    row["inst"] = v
                v = _col(cols, os_cpu_idx)
                if v is not None:
                    row["os_cpu"] = float(v)
                v = _col(cols, os_cpu_max_idx)
                if v is not None:
                    row["os_cpu_max"] = float(v)
                v = _col(cols, cpu_per_s_idx)
                if v is not None:
                    row["cpu_per_s"] = float(v)
                v = _col(cols, cpu_per_s_max_idx)
                if v is not None:
                    row["cpu_per_s_max"] = float(v)
                v = _col(cols, read_iops_idx)
                if v is not None:
                    row["read_iops"] = float(v)
                v = _col(cols, write_iops_idx)
                if v is not None:
                    row["write_iops"] = float(v)
                v = _col(cols, read_iops_max_idx)
                if v is not None:
                    row["read_iops_max"] = float(v)
                v = _col(cols, write_iops_max_idx)
                if v is not None:
                    row["write_iops_max"] = float(v)
                v = _col(cols, redo_idx)
                if v is not None:
                    row["redo_mb_s"] = float(v)
                v = _col(cols, dur_idx)
                if v is not None:
                    row["dur_m"] = float(v)
                rows.append(row)
            except (ValueError, IndexError):
                continue

        if not rows:
            return {}

        result: dict = {}

        # 스냅샷별로 그룹핑 → 인스턴스 합산(IOPS, CPU/s) 또는 평균(CPU%)
        snap_groups: dict[str, list[dict]] = {}
        for row in rows:
            snap_key = row.get("snap", "0")
            snap_groups.setdefault(snap_key, []).append(row)

        # 스냅샷별 집계값 계산
        snap_os_cpu: list[float] = []
        snap_os_cpu_max: list[float] = []
        snap_cpu_per_s: list[float] = []
        snap_cpu_per_s_max: list[float] = []
        snap_iops: list[float] = []
        snap_iops_max: list[float] = []
        snap_redo: list[float] = []
        snap_dur: list[float] = []

        for snap_key, snap_rows in snap_groups.items():
            # CPU%: 인스턴스 평균 (각 인스턴스의 호스트 CPU 사용률)
            cpu_vals = [r["os_cpu"] for r in snap_rows if "os_cpu" in r]
            if cpu_vals:
                snap_os_cpu.append(sum(cpu_vals) / len(cpu_vals))

            cpu_max_vals = [r["os_cpu_max"] for r in snap_rows if "os_cpu_max" in r]
            if cpu_max_vals:
                snap_os_cpu_max.append(max(cpu_max_vals))

            # CPU/s: 인스턴스 합산 (RAC 전체 DB CPU 사용량)
            cps_vals = [r["cpu_per_s"] for r in snap_rows if "cpu_per_s" in r]
            if cps_vals:
                snap_cpu_per_s.append(sum(cps_vals))

            cps_max_vals = [r["cpu_per_s_max"] for r in snap_rows if "cpu_per_s_max" in r]
            if cps_max_vals:
                snap_cpu_per_s_max.append(sum(cps_max_vals))

            # IOPS: 인스턴스 합산
            iops_vals = []
            for r in snap_rows:
                ri = r.get("read_iops", 0)
                wi = r.get("write_iops", 0)
                iops_vals.append(ri + wi)
            if iops_vals:
                snap_iops.append(sum(iops_vals))

            iops_max_vals = []
            for r in snap_rows:
                ri = r.get("read_iops_max", 0)
                wi = r.get("write_iops_max", 0)
                iops_max_vals.append(ri + wi)
            if iops_max_vals:
                snap_iops_max.append(sum(iops_max_vals))

            # Redo: 인스턴스 합산
            redo_vals = [r["redo_mb_s"] for r in snap_rows if "redo_mb_s" in r]
            if redo_vals:
                snap_redo.append(sum(redo_vals))

            # dur_m: 첫 번째 값 사용
            dur_vals = [r["dur_m"] for r in snap_rows if "dur_m" in r]
            if dur_vals:
                snap_dur.append(dur_vals[0])

        # 스냅샷 평균 → 최종 결과
        if snap_os_cpu:
            result["avg_cpu_percent"] = round(sum(snap_os_cpu) / len(snap_os_cpu), 1)
        if snap_os_cpu_max:
            result["peak_cpu_percent"] = round(max(snap_os_cpu_max), 1)
        if snap_cpu_per_s:
            result["avg_cpu_per_s"] = round(sum(snap_cpu_per_s) / len(snap_cpu_per_s), 3)
        if snap_cpu_per_s_max:
            result["peak_cpu_per_s"] = round(max(snap_cpu_per_s_max), 3)
        if snap_iops:
            result["avg_iops"] = round(sum(snap_iops) / len(snap_iops), 0)
        if snap_iops_max:
            result["peak_iops"] = round(max(snap_iops_max), 0)

        # Redo → 일별 바이트
        if snap_redo:
            avg_redo_mb_s = sum(snap_redo) / len(snap_redo)
            result["redo_bytes_per_day"] = avg_redo_mb_s * 86400 * 1024 * 1024

        # dur_m (스냅샷 기간) 저장
        if snap_dur:
            result["dur_m"] = sum(snap_dur) / len(snap_dur)

        return result

    def _parse_memory_section(self, content: str) -> dict:
        """AWR .out의 MEMORY 섹션에서 메모리 사용량을 파싱합니다.

        RAC 환경에서는 인스턴스별 TOTAL(SGA+PGA)을 합산합니다.
        """
        import re

        match = re.search(
            r"~~BEGIN-MEMORY~~\s*\n(.*?)~~END-MEMORY~~",
            content, re.DOTALL,
        )
        if not match:
            return {}

        section = match.group(1)
        lines = section.strip().split("\n")

        # 헤더 행 찾기
        header_line = None
        data_start = 0
        for i, line in enumerate(lines):
            stripped = line.strip()
            if "SNAP_ID" in stripped or "TOTAL" in stripped:
                header_line = stripped
                data_start = i + 1
                break

        if not header_line:
            return {}

        headers = header_line.split()
        col_map = {h: idx for idx, h in enumerate(headers)}

        snap_idx = col_map.get("SNAP_ID")
        inst_idx = col_map.get("INSTANCE_NUMBER")
        total_idx = col_map.get("TOTAL")

        if total_idx is None:
            return {}

        # 구분선 건너뛰기
        for i in range(data_start, len(lines)):
            if lines[i].strip().startswith("---"):
                data_start = i + 1
                break

        # 데이터 행 파싱
        rows: list[dict] = []
        for i in range(data_start, len(lines)):
            line = lines[i].strip()
            if not line or line.startswith("~~"):
                break
            cols = line.split()
            try:
                row: dict = {}
                if snap_idx is not None and len(cols) > snap_idx:
                    row["snap"] = cols[snap_idx]
                if inst_idx is not None and len(cols) > inst_idx:
                    row["inst"] = cols[inst_idx]
                if total_idx is not None and len(cols) > total_idx:
                    row["total"] = float(cols[total_idx])
                rows.append(row)
            except (ValueError, IndexError):
                continue

        if not rows:
            return {}

        # 스냅샷별 인스턴스 합산
        snap_totals: dict[str, float] = {}
        for row in rows:
            snap_key = row.get("snap", "0")
            snap_totals[snap_key] = snap_totals.get(snap_key, 0) + row.get("total", 0)

        totals = list(snap_totals.values())
        if not totals:
            return {}

        return {
            "avg_memory_gb": round(sum(totals) / len(totals), 1),
            "peak_memory_gb": round(max(totals), 1),
        }

    def _parse_sga_advice(self, content: str) -> dict:
        """AWR .out의 SGA-ADVICE 섹션에서 SGA 분석 결과를 파싱합니다.

        현재 SGA 크기(SGA_SIZE_FACTOR=1.00)와 권장 SGA 크기를 추출합니다.
        권장 SGA: ESTD_DB_TIME_FACTOR가 0.90 이하인 최소 SGA 크기.
        """
        import re

        match = re.search(
            r"~~BEGIN-SGA-ADVICE~~\s*\n(.*?)~~END-SGA-ADVICE~~",
            content, re.DOTALL,
        )
        if not match:
            return {}

        section = match.group(1)
        lines = section.strip().split("\n")

        # 헤더 행 찾기
        header_line = None
        data_start = 0
        for i, line in enumerate(lines):
            stripped = line.strip()
            if "SGA_SIZE" in stripped or "INST_ID" in stripped:
                header_line = stripped
                data_start = i + 1
                break

        if not header_line:
            return {}

        headers = header_line.split()
        col_map = {h: idx for idx, h in enumerate(headers)}

        inst_idx = col_map.get("INST_ID")
        sga_size_idx = col_map.get("SGA_SIZE")
        factor_idx = col_map.get("SGA_SIZE_FACTOR")
        db_time_factor_idx = col_map.get("ESTD_DB_TIME_FACTOR")

        if sga_size_idx is None or factor_idx is None:
            return {}

        # 구분선 건너뛰기
        for i in range(data_start, len(lines)):
            if lines[i].strip().startswith("---"):
                data_start = i + 1
                break

        # 데이터 행 파싱 (인스턴스 1만 사용)
        current_sga: float | None = None
        recommended_sga: float | None = None

        for i in range(data_start, len(lines)):
            line = lines[i].strip()
            if not line or line.startswith("~~"):
                break
            cols = line.split()
            try:
                inst = int(cols[inst_idx]) if inst_idx is not None else 1
                if inst != 1:
                    continue

                sga_size = float(cols[sga_size_idx])
                sga_factor = float(cols[factor_idx])

                # 현재 SGA: factor = 1.00
                if abs(sga_factor - 1.0) < 0.01:
                    current_sga = sga_size

                # 권장 SGA: DB Time Factor ≤ 0.90인 최소 크기
                if db_time_factor_idx is not None and len(cols) > db_time_factor_idx:
                    db_time_factor = float(cols[db_time_factor_idx])
                    if db_time_factor <= 0.90:
                        if recommended_sga is None or sga_size < recommended_sga:
                            recommended_sga = sga_size
            except (ValueError, IndexError):
                continue

        result: dict = {}
        if current_sga is not None:
            result["current_sga_gb"] = current_sga
        if recommended_sga is not None:
            result["recommended_sga_gb"] = recommended_sga

        return result

    def _parse_awr_out_network_from_content(self, content: str) -> dict:
        """AWR .out 파일 내용에서 네트워크/Redo 데이터를 파싱합니다.

        기존 _parse_awr_out_network의 내부 로직을 content 기반으로 분리.
        """
        result: dict = {}

        # MAIN-METRICS에서 dur_m 확보
        main_metrics = self._parse_main_metrics_section(content)

        # SYSSTAT 섹션 파싱
        sysstat_data = self._parse_sysstat_section(content)
        if sysstat_data:
            dur_m = (main_metrics or {}).get("dur_m") or sysstat_data.get("dur_m", 60)
            factor = (1440.0 / dur_m) * 1024 * 1024

            if sysstat_data.get("network_outgoing_mb") is not None:
                result["sent_bytes_per_day"] = sysstat_data["network_outgoing_mb"] * factor

            if sysstat_data.get("network_incoming_mb") is not None:
                result["recv_bytes_per_day"] = sysstat_data["network_incoming_mb"] * factor

        # Redo는 _parse_main_metrics_full에서 이미 처리되므로 여기서는 스킵
        # (중복 방지)

        return result

    def _supplement_awr_network_data(
        self, file_path: str, parsed: ParsedDocumentInfo
    ) -> None:
        """AWR .out 파일에서 네트워크/Redo 데이터를 직접 파싱합니다.

        Bedrock(AI)은 AWR .out 파일의 넓은 테이블에서 네트워크 컬럼을
        정확히 추출하지 못하는 경우가 많으므로, 직접 파싱 결과를 우선 사용합니다.
        직접 파싱에 성공하면 Bedrock 결과를 덮어씁니다.

        파싱 대상:
        - SYSSTAT 섹션: network_incoming_mb, network_outgoing_mb
        - MAIN-METRICS 섹션: redo_mb_s, dur_m (스냅샷 기간)
        """
        # AWR .out 파일 찾기
        out_files = self._find_awr_out_files(file_path)
        if not out_files:
            logger.debug("AWR .out 파일 없음, 네트워크 데이터 직접 파싱 스킵")
            return

        awr = parsed.awr_metrics

        for out_file in out_files:
            logger.info("AWR .out 파일 직접 파싱 (AI 대신 우선 사용): %s", out_file)
            net_data = self._parse_awr_out_network(out_file)
            if not net_data:
                continue

            # 직접 파싱 결과로 덮어쓰기 (Bedrock 결과보다 정확)
            if net_data.get("sent_bytes_per_day") is not None:
                awr.sqlnet_bytes_sent_per_day = net_data["sent_bytes_per_day"]
                logger.info("네트워크 송신 (직접 파싱): %.0f bytes/일", awr.sqlnet_bytes_sent_per_day)

            if net_data.get("recv_bytes_per_day") is not None:
                awr.sqlnet_bytes_received_per_day = net_data["recv_bytes_per_day"]
                logger.info("네트워크 수신 (직접 파싱): %.0f bytes/일", awr.sqlnet_bytes_received_per_day)

            if net_data.get("redo_bytes_per_day") is not None:
                awr.redo_bytes_per_day = net_data["redo_bytes_per_day"]
                logger.info("Redo 생성량 (직접 파싱): %.0f bytes/일", awr.redo_bytes_per_day)


    def _find_awr_out_files(self, file_path: str) -> list[str]:
        """AWR .out 파일을 찾습니다."""
        if os.path.isfile(file_path) and file_path.endswith(".out"):
            return [file_path]

        if os.path.isdir(file_path):
            out_files = []
            for entry in sorted(os.listdir(file_path)):
                if entry.endswith(".out"):
                    out_files.append(os.path.join(file_path, entry))
            return out_files

        return []

    def _parse_awr_out_network(self, out_file: str) -> dict:
        """AWR .out 파일에서 네트워크/Redo 데이터를 파싱합니다.

        SYSSTAT 섹션에서 network_incoming_mb/network_outgoing_mb를,
        MAIN-METRICS 섹션에서 redo_mb_s와 dur_m을 추출합니다.

        Returns:
            dict with keys: sent_bytes_per_day, recv_bytes_per_day, redo_bytes_per_day
        """
        import re

        try:
            with open(out_file, "r", encoding="utf-8") as f:
                content = f.read()
        except Exception as e:
            logger.warning("AWR .out 파일 읽기 실패: %s - %s", out_file, e)
            return {}

        result: dict = {}

        # MAIN-METRICS 섹션을 먼저 파싱하여 실제 dur_m 확보
        main_metrics = self._parse_main_metrics_section(content)

        # SYSSTAT 섹션 파싱: network_incoming_mb, network_outgoing_mb
        sysstat_data = self._parse_sysstat_section(content)
        if sysstat_data:
            # SYSSTAT의 값은 스냅샷 기간(보통 60분) 동안의 MB 값
            # MAIN-METRICS에서 실제 dur_m을 가져와 사용
            dur_m = (main_metrics or {}).get("dur_m") or sysstat_data.get("dur_m", 60)
            # 일별로 변환: MB × (1440 / dur_m) × 1024 × 1024
            factor = (1440.0 / dur_m) * 1024 * 1024

            if sysstat_data.get("network_outgoing_mb") is not None:
                result["sent_bytes_per_day"] = sysstat_data["network_outgoing_mb"] * factor

            if sysstat_data.get("network_incoming_mb") is not None:
                result["recv_bytes_per_day"] = sysstat_data["network_incoming_mb"] * factor

        # MAIN-METRICS 섹션에서 redo_mb_s 추출 (이미 위에서 파싱됨)
        if main_metrics and main_metrics.get("redo_mb_s") is not None:
            # redo_mb_s는 초당 MB → 일별 바이트: × 86400 × 1024 × 1024
            result["redo_bytes_per_day"] = (
                main_metrics["redo_mb_s"] * 86400 * 1024 * 1024
            )

        return result

    def _parse_sysstat_section(self, content: str) -> dict:
        """AWR .out의 SYSSTAT 섹션에서 네트워크 데이터를 파싱합니다.

        여러 스냅샷의 평균값을 반환합니다.
        """
        import re

        # SYSSTAT 섹션 추출
        match = re.search(
            r"~~BEGIN-SYSSTAT~~\s*\n(.*?)~~END-SYSSTAT~~",
            content, re.DOTALL
        )
        if not match:
            return {}

        section = match.group(1)
        lines = section.strip().split("\n")

        # 헤더 행 찾기
        header_line = None
        data_start = 0
        for i, line in enumerate(lines):
            stripped = line.strip()
            if stripped.startswith("SNAP_ID") or "network_incoming_mb" in stripped:
                header_line = stripped
                data_start = i + 1
                break

        if not header_line:
            return {}

        # 헤더에서 컬럼 인덱스 찾기
        headers = header_line.split()
        col_map = {h: idx for idx, h in enumerate(headers)}

        incoming_idx = col_map.get("network_incoming_mb")
        outgoing_idx = col_map.get("network_outgoing_mb")

        if incoming_idx is None and outgoing_idx is None:
            return {}

        # 구분선 건너뛰기
        for i in range(data_start, len(lines)):
            if lines[i].strip().startswith("---"):
                data_start = i + 1
                break

        # 데이터 행 파싱
        incoming_vals: list[float] = []
        outgoing_vals: list[float] = []

        for i in range(data_start, len(lines)):
            line = lines[i].strip()
            if not line or line.startswith("~~"):
                break
            cols = line.split()
            if len(cols) <= max(incoming_idx or 0, outgoing_idx or 0):
                continue
            try:
                if incoming_idx is not None:
                    incoming_vals.append(float(cols[incoming_idx]))
                if outgoing_idx is not None:
                    outgoing_vals.append(float(cols[outgoing_idx]))
            except (ValueError, IndexError):
                continue

        result: dict = {}
        if incoming_vals:
            result["network_incoming_mb"] = sum(incoming_vals) / len(incoming_vals)
        if outgoing_vals:
            result["network_outgoing_mb"] = sum(outgoing_vals) / len(outgoing_vals)

        # dur_m은 MAIN-METRICS에서 가져오지만, 기본값 60분 사용
        result["dur_m"] = 60

        return result

    def _parse_main_metrics_section(self, content: str) -> dict:
        """AWR .out의 MAIN-METRICS 섹션에서 redo_mb_s와 dur_m을 파싱합니다.

        여러 스냅샷/인스턴스의 평균값을 반환합니다.
        'end' 컬럼의 날짜+시간 공백 오프셋을 보정합니다.
        """
        import re

        # MAIN-METRICS 섹션 추출
        match = re.search(
            r"~~BEGIN-MAIN-METRICS~~\s*\n(.*?)~~END-MAIN-METRICS~~",
            content, re.DOTALL
        )
        if not match:
            return {}

        section = match.group(1)
        lines = section.strip().split("\n")

        # 헤더 행 찾기
        header_line = None
        data_start = 0
        for i, line in enumerate(lines):
            stripped = line.strip()
            if stripped.startswith("snap") or "redo_mb_s" in stripped:
                header_line = stripped
                data_start = i + 1
                break

        if not header_line:
            return {}

        # 헤더에서 컬럼 인덱스 찾기
        headers = header_line.split()
        col_map = {h: idx for idx, h in enumerate(headers)}

        # 'end' 컬럼 위치 (날짜+시간이 2개 토큰으로 분리됨)
        end_header_idx = col_map.get("end")

        redo_idx = col_map.get("redo_mb_s")
        dur_idx = col_map.get("dur_m")

        if redo_idx is None:
            return {}

        def _actual_idx(header_idx: int) -> int:
            """end 컬럼 이후만 +1 오프셋 적용."""
            if end_header_idx is not None and header_idx > end_header_idx:
                return header_idx + 1
            return header_idx

        redo_actual = _actual_idx(redo_idx)
        dur_actual = _actual_idx(dur_idx) if dur_idx is not None else None

        # 구분선 건너뛰기
        for i in range(data_start, len(lines)):
            if lines[i].strip().startswith("---"):
                data_start = i + 1
                break

        # 데이터 행 파싱
        redo_vals: list[float] = []
        dur_vals: list[float] = []

        for i in range(data_start, len(lines)):
            line = lines[i].strip()
            if not line or line.startswith("~~"):
                break
            cols = line.split()
            if len(cols) <= redo_actual:
                continue
            try:
                redo_vals.append(float(cols[redo_actual]))
                if dur_actual is not None and len(cols) > dur_actual:
                    dur_vals.append(float(cols[dur_actual]))
            except (ValueError, IndexError):
                continue

        result: dict = {}
        if redo_vals:
            result["redo_mb_s"] = sum(redo_vals) / len(redo_vals)
        if dur_vals:
            result["dur_m"] = sum(dur_vals) / len(dur_vals)

        return result

    # 타겟 엔진 키워드 → 엔진 코드 매핑
    _TARGET_ENGINE_MAP: dict[str, str] = {
        "aurora postgresql": "aurora-postgresql",
        "aurora postgres": "aurora-postgresql",
        "aurora mysql": "aurora-mysql",
        "rds for postgresql": "postgresql",
        "rds postgresql": "postgresql",
        "rds for mysql": "mysql",
        "rds mysql": "mysql",
        "rds for sql server": "sqlserver-ee",
        "rds sql server": "sqlserver-ee",
        # Oracle SE2 (LI)를 먼저 매칭해야 "rds for oracle"보다 우선 적용됨
        "rds for oracle se2": "oracle-se2",
        "oracle se2": "oracle-se2",
        "rds for oracle ee": "oracle-ee",
        "oracle ee": "oracle-ee",
        "rds for oracle": "oracle-ee",
    }

    def _supplement_target_engine(
        self, file_path: str, parsed: "ParsedDocumentInfo"
    ) -> None:
        """migration_recommendation.md에서 '추천 타겟' 필드를 파싱하여 target_engine을 보완합니다.

        Bedrock이 소스 엔진(oracle-ee)을 반환하는 경우가 많으므로,
        migration_recommendation.md의 '추천 타겟' 텍스트에서 타겟 엔진을 직접 추출합니다.
        """
        if parsed.target_engine is not None:
            return

        # migration_recommendation.md 파일 찾기
        rec_file = self._find_migration_recommendation(file_path)
        if not rec_file:
            return

        try:
            with open(rec_file, "r", encoding="utf-8") as f:
                content = f.read()
        except Exception as e:
            logger.warning("migration_recommendation.md 읽기 실패: %s", e)
            return

        # "**추천 타겟**: Aurora PostgreSQL" 패턴 매칭
        import re
        match = re.search(
            r"\*\*추천\s*타겟\*\*\s*:\s*(.+)",
            content,
        )
        if not match:
            return

        target_text = match.group(1).strip().lower()
        for keyword, engine_code in self._TARGET_ENGINE_MAP.items():
            if keyword in target_text:
                parsed.target_engine = engine_code
                logger.info(
                    "migration_recommendation.md에서 타겟 엔진 추출: %s → %s",
                    match.group(1).strip(), engine_code,
                )
                return

    def _find_migration_recommendation(self, file_path: str) -> str | None:
        """입력 경로에서 migration_recommendation.md 파일을 찾습니다."""
        if os.path.isdir(file_path):
            candidate = os.path.join(file_path, "migration_recommendation.md")
            if os.path.isfile(candidate):
                return candidate
        elif os.path.isfile(file_path):
            # 같은 디렉토리에서 찾기
            parent = os.path.dirname(file_path)
            candidate = os.path.join(parent, "migration_recommendation.md")
            if os.path.isfile(candidate):
                return candidate
        return None

